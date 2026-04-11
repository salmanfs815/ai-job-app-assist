from io import BytesIO
from pathlib import Path
from typing import Literal

from docx import Document
from pypdf import PdfReader
import pypdfium2 as pdfium
import pytesseract
from PIL import Image

from app.config import settings

ALLOWED_EXTENSIONS = {".pdf", ".docx"}
OcrStatus = Literal["not_used", "used", "failed"]


def extract_text_from_resume_file(filename: str, content: bytes) -> tuple[str, OcrStatus]:
    ext = Path(filename).suffix.lower()
    if ext not in ALLOWED_EXTENSIONS:
        raise ValueError("Unsupported resume format. Please upload a PDF or DOCX file.")

    if ext == ".pdf":
        return _extract_pdf_text(content)

    if ext == ".docx":
        return _extract_docx_text(content), "not_used"

    raise ValueError("Unsupported resume format.")


def _extract_pdf_text(content: bytes) -> tuple[str, OcrStatus]:
    reader = PdfReader(BytesIO(content))
    text_parts: list[str] = []
    for page in reader.pages:
        text_parts.append(page.extract_text() or "")
    extracted = "\n".join(text_parts).strip()
    if extracted:
        return extracted, "not_used"
    ocr_text = _extract_pdf_text_with_ocr(content)
    if ocr_text:
        return ocr_text, "used"
    return "", "failed"


def _extract_docx_text(content: bytes) -> str:
    doc = Document(BytesIO(content))
    paragraphs = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
    return "\n".join(paragraphs).strip()


def _extract_pdf_text_with_ocr(content: bytes) -> str:
    if not settings.enable_pdf_ocr_fallback:
        return ""

    if settings.tesseract_cmd.strip():
        pytesseract.pytesseract.tesseract_cmd = settings.tesseract_cmd.strip()

    try:
        pdf = pdfium.PdfDocument(BytesIO(content))
        text_chunks: list[str] = []

        for i in range(len(pdf)):
            page = pdf[i]
            bitmap = page.render(scale=2.0)
            pil_image: Image.Image = bitmap.to_pil()
            chunk = pytesseract.image_to_string(pil_image).strip()
            if chunk:
                text_chunks.append(chunk)

        return "\n".join(text_chunks).strip()
    except Exception:
        # OCR is best-effort; return empty so caller can raise a friendly extraction message.
        return ""
